"""
Document Processor Module

Extracts text from PDF and DOCX files for translation.
"""

import os
import sys
from typing import Optional


def detect_columns(words, page_width, min_gap_ratio=0.05):
    """
    Detect column boundaries based on word positions.
    Returns list of (left_bound, right_bound) tuples for each detected column.
    """
    if not words:
        return [(0, page_width)]
    
    # Get all x0 (left edge) positions
    x_positions = sorted(set(w['x0'] for w in words))
    
    if len(x_positions) < 2:
        return [(0, page_width)]
    
    # Find gaps between words that indicate column boundaries
    min_gap = page_width * min_gap_ratio  # Minimum gap to consider as column separator
    
    # Group words by their approximate x position to find column clusters
    x0_values = [w['x0'] for w in words]
    x1_values = [w['x1'] for w in words]
    
    # Find large horizontal gaps that indicate column separations
    gaps = []
    sorted_by_x = sorted(words, key=lambda w: w['x0'])
    
    for i in range(len(sorted_by_x) - 1):
        current_right = sorted_by_x[i]['x1']
        next_left = sorted_by_x[i + 1]['x0']
        gap = next_left - current_right
        
        if gap > min_gap:
            gaps.append((current_right, next_left, gap))
    
    # If no significant gaps found, treat as single column
    if not gaps:
        return [(min(x0_values), max(x1_values))]
    
    # Sort gaps by size and take the largest ones (likely column separators)
    gaps.sort(key=lambda g: g[2], reverse=True)
    
    # Use top gaps to define column boundaries (limit to 3 columns max for most documents)
    column_separators = sorted([g[0] + (g[1] - g[0]) / 2 for g in gaps[:2]])
    
    # Create column boundaries
    columns = []
    prev_boundary = 0
    for sep in column_separators:
        columns.append((prev_boundary, sep))
        prev_boundary = sep
    columns.append((prev_boundary, page_width))
    
    return columns


def extract_text_by_columns(page, columns):
    """
    Extract text from page respecting column boundaries.
    Reads each column top-to-bottom, then moves to next column.
    Each word is assigned to exactly ONE column (the best match) to prevent duplicates.
    """
    words = page.extract_words(x_tolerance=3, y_tolerance=3)
    
    if not words:
        return ""
    
    # Assign each word to exactly ONE column (the best-matching one)
    # This prevents words near column boundaries from appearing in multiple columns
    column_word_map = {i: [] for i in range(len(columns))}
    
    for word in words:
        word_center = (word['x0'] + word['x1']) / 2
        best_col_idx = 0
        best_distance = float('inf')
        
        for col_idx, (col_left, col_right) in enumerate(columns):
            col_center = (col_left + col_right) / 2
            distance = abs(word_center - col_center)
            if distance < best_distance:
                best_distance = distance
                best_col_idx = col_idx
        
        column_word_map[best_col_idx].append(word)
    
    column_texts = []
    
    for col_idx in range(len(columns)):
        col_words = column_word_map[col_idx]
        
        if not col_words:
            continue
        
        # Sort by y position (top to bottom), then x position (left to right)
        col_words.sort(key=lambda w: (w['top'], w['x0']))
        
        # Group words into lines based on y position
        lines = []
        current_line = []
        current_y = None
        y_tolerance = 5
        
        for word in col_words:
            if current_y is None or abs(word['top'] - current_y) <= y_tolerance:
                current_line.append(word)
                current_y = word['top'] if current_y is None else current_y
            else:
                if current_line:
                    # Sort words in line by x position
                    current_line.sort(key=lambda w: w['x0'])
                    lines.append(' '.join(w['text'] for w in current_line))
                current_line = [word]
                current_y = word['top']
        
        # Don't forget the last line
        if current_line:
            current_line.sort(key=lambda w: w['x0'])
            lines.append(' '.join(w['text'] for w in current_line))
        
        if lines:
            column_texts.append('\n'.join(lines))
    
    return '\n\n---\n\n'.join(column_texts)


def extract_text_from_pdf(file_path: str) -> dict:
    """
    Extract text from a PDF file with intelligent multi-column detection.
    
    This improved version:
    1. Detects multi-column layouts automatically
    2. Reads each column top-to-bottom before moving to the next
    3. Preserves reading order for complex documents like planning maps
    
    Args:
        file_path: Path to the PDF file
    
    Returns:
        dict with success status and extracted text
    """
    try:
        import pdfplumber
        import re
        
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_width = page.width
                page_height = page.height
                
                # Extract words with positions
                words = page.extract_words(x_tolerance=3, y_tolerance=3)
                
                if words:
                    # Detect columns based on word positions
                    columns = detect_columns(words, page_width)
                    num_columns = len(columns)
                    
                    print(f"DEBUG: Page {page_num} - Detected {num_columns} column(s)", file=sys.stderr)
                    
                    if num_columns > 1:
                        # Multi-column layout - extract by columns
                        page_text = extract_text_by_columns(page, columns)
                    else:
                        # Single column - use standard extraction
                        page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                    
                    if page_text:
                        # Cleanup
                        clean_text = re.sub(r'[ ]{2,}', ' ', page_text)
                        clean_text = re.sub(r'(\w+)-\n(\w+)', r'\1\2', clean_text)
                        text_content.append(f"--- Page {page_num} ---\n{clean_text}")
                else:
                    # No words found, try standard extraction
                    page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                    if page_text:
                        clean_text = re.sub(r'[ ]{2,}', ' ', page_text)
                        clean_text = re.sub(r'(\w+)-\n(\w+)', r'\1\2', clean_text)
                        text_content.append(f"--- Page {page_num} ---\n{clean_text}")
        
        full_text = "\n\n".join(text_content)
        
        if not full_text.strip():
             print(f"DEBUG: No text found via native extraction. Attempting Vision OCR...", file=sys.stderr)
             return ocr_pdf_with_vision(file_path)
        
        return {
            "success": True,
            "text": full_text,
            "page_count": len(text_content),
            "file_type": "pdf",
            "error": None
        }
        
    except Exception as e:
        return {
            "success": False,
            "text": None,
            "error": f"Failed to extract PDF text: {str(e)}"
        }

def ocr_pdf_with_vision(file_path: str) -> dict:
    """
    Extract text from PDF using OCR (Azure OpenAI Vision) via PyMuPDF rendering.
    Note: This feature requires PyMuPDF which is not available on Render free tier.
    """
    try:
        import fitz  # PyMuPDF
        from openai_client import extract_text_from_image
        
        doc = fitz.open(file_path)
        text_parts = []
        
        print(f"DEBUG: Starting OCR for {len(doc)} pages...", file=sys.stderr)
        
        for i, page in enumerate(doc):
            # Render page to image (scale 2.0 for better quality)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_bytes = pix.tobytes("png")
            
            result = extract_text_from_image(img_bytes)
            
            if result["success"]:
                text_parts.append(f"--- Page {i+1} ---\n{result['text']}")
                print(f"DEBUG: OCR Page {i+1} success", file=sys.stderr)
            else:
                text_parts.append(f"--- Page {i+1} ---\n[OCR Failed]")
                print(f"DEBUG: OCR Page {i+1} failed: {result.get('error')}", file=sys.stderr)
        
        doc.close()
        full_text = "\n\n".join(text_parts)
        
        return {
            "success": True,
            "text": full_text,
            "page_count": len(text_parts),
            "file_type": "pdf_ocr",
            "error": None
        }
        
    except ImportError:
        # PyMuPDF not available (e.g., on Render deployment)
        print(f"DEBUG: PyMuPDF not available, OCR disabled", file=sys.stderr)
        return {
            "success": False,
            "text": None,
            "error": "OCR feature not available on this deployment (PyMuPDF not installed)"
        }
    except Exception as e:
        print(f"DEBUG: OCR failed: {str(e)}", file=sys.stderr)
        return {"success": False, "text": None, "error": f"OCR failed: {str(e)}"}


def extract_text_from_docx(file_path: str) -> dict:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
    
    Returns:
        dict with success status and extracted text
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n\n".join(paragraphs)
        
        return {
            "success": True,
            "text": full_text,
            "paragraph_count": len(paragraphs),
            "file_type": "docx",
            "error": None
        }
        
    except Exception as e:
        return {
            "success": False,
            "text": None,
            "error": f"Failed to extract DOCX text: {str(e)}"
        }


def extract_text_from_txt(file_path: str) -> dict:
    """
    Read text from a plain text file.
    
    Args:
        file_path: Path to the text file
    
    Returns:
        dict with success status and text content
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        return {
            "success": True,
            "text": text,
            "file_type": "txt",
            "error": None
        }
        
    except Exception as e:
        return {
            "success": False,
            "text": None,
            "error": f"Failed to read text file: {str(e)}"
        }


def extract_text(file_path: str, file_type: Optional[str] = None) -> dict:
    """
    Extract text from a document based on file type.
    
    Args:
        file_path: Path to the document
        file_type: Optional file type override ('pdf', 'docx', 'txt')
    
    Returns:
        dict with extracted text or error
    """
    if not os.path.exists(file_path):
        return {
            "success": False,
            "text": None,
            "error": f"File not found: {file_path}"
        }
    
    # Determine file type from extension if not provided
    if not file_type:
        _, ext = os.path.splitext(file_path)
        file_type = ext.lower().lstrip('.')
    
    # Route to appropriate extractor
    extractors = {
        'pdf': extract_text_from_pdf,
        'docx': extract_text_from_docx,
        'txt': extract_text_from_txt,
        'text': extract_text_from_txt,
    }
    
    extractor = extractors.get(file_type)
    
    if not extractor:
        return {
            "success": False,
            "text": None,
            "error": f"Unsupported file type: {file_type}. Supported: pdf, docx, txt"
        }
    
    return extractor(file_path)


# ... (existing imports)
from translator import translate_text, translate_long_text

def translate_paragraph_runs(para, source_lang: str, target_lang: str) -> int:
    """
    Translate paragraph by translating each run individually to preserve formatting.
    Returns count of translated runs.
    """
    translated_count = 0
    runs_with_text = [r for r in para.runs if r.text and r.text.strip()]
    
    if not runs_with_text:
        # No runs with text - paragraph might have text but no runs (common in converted PDFs)
        return 0
    
    for run in runs_with_text:
        # Use translate_long_text for potentially large runs
        result = translate_long_text(run.text, source_lang, target_lang)
        if result["success"]:
            # Preserve the run's text while keeping all formatting (bold, italic, font, etc.)
            run.text = result["translated_text"]
            translated_count += 1
        else:
            print(f"DEBUG: Run translation failed: {result.get('error')}", file=sys.stderr)
    
    return translated_count

def translate_docx_file(file_path: str, output_path: str, source_lang: str, target_lang: str) -> dict:
    """
    Translate a DOCX file preserving formatting.
    """
    try:
        from docx import Document
        
        print(f"DEBUG: Opening DOCX: {file_path}", file=sys.stderr)
        doc = Document(file_path)
        
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        print(f"DEBUG: Found {para_count} paragraphs and {table_count} tables.", file=sys.stderr)
        
        translated_count = 0
        
        # Track already-translated paragraphs to prevent duplicates
        # Use the XML element itself (not id()) because python-docx may create
        # different wrapper objects for the same underlying XML element
        translated_elements = set()
        
        # Also track translated text content to catch duplicate text across different elements
        translated_texts = set()
        
        # Translate paragraphs - use run-level translation to preserve formatting
        for i, para in enumerate(doc.paragraphs):
            para_elem = para._element
            if para_elem in translated_elements:
                print(f"DEBUG: Para {i} - SKIPPING (already translated, same XML element)", file=sys.stderr)
                continue
            para_text = para.text.strip()
            if para_text:
                # Skip if we've already translated identical text content
                if para_text in translated_texts:
                    print(f"DEBUG: Para {i} - SKIPPING (duplicate text content)", file=sys.stderr)
                    # Still translate it since it's a different element, but note it
                
                # Translate each run individually to preserve bold, italic, fonts, etc.
                runs_translated = translate_paragraph_runs(para, source_lang, target_lang)
                if runs_translated > 0:
                    translated_count += runs_translated
                    translated_elements.add(para_elem)
                    translated_texts.add(para_text)
                    print(f"DEBUG: Para {i} - translated {runs_translated} runs", file=sys.stderr)
                elif para_text:
                    # Fallback: If no runs but text exists, translate whole paragraph
                    print(f"DEBUG: Para {i} - no runs, using paragraph-level translation ({len(para_text)} chars)", file=sys.stderr)
                    result = translate_long_text(para_text, source_lang, target_lang)
                    if result["success"]:
                        para.text = result["translated_text"]
                        translated_count += 1
                        translated_elements.add(para_elem)
                        translated_texts.add(para_text)
                    else:
                        print(f"DEBUG: Translation failed for para {i}: {result.get('error')}", file=sys.stderr)

        # Translate tables - use run-level translation to preserve formatting
        # Track visited cells using XML element identity to avoid translating merged cells multiple times
        visited_cell_elements = set()
        for t_idx, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    cell_elem = cell._element
                    if cell_elem in visited_cell_elements:
                        continue
                    visited_cell_elements.add(cell_elem)
                    for para in cell.paragraphs:
                        para_elem = para._element
                        if para_elem in translated_elements:
                            continue
                        para_text = para.text.strip()
                        if para_text:
                            # Translate each run individually to preserve formatting
                            runs_translated = translate_paragraph_runs(para, source_lang, target_lang)
                            if runs_translated > 0:
                                translated_count += runs_translated
                                translated_elements.add(para_elem)
                                translated_texts.add(para_text)
                            elif para_text:
                                # Fallback for paragraphs without runs
                                result = translate_long_text(para_text, source_lang, target_lang)
                                if result["success"]:
                                    para.text = result["translated_text"]
                                    translated_count += 1
                                    translated_elements.add(para_elem)
                                    translated_texts.add(para_text)
        
        print(f"DEBUG: Total translated segments: {translated_count}", file=sys.stderr)
        
        if translated_count == 0 and (para_count > 0 or table_count > 0):
             print(f"WARNING: Content found but nothing translated (maybe empty text?)", file=sys.stderr)
        elif para_count == 0 and table_count == 0:
             print(f"WARNING: No paragraphs or tables found in DOCX! (Possible scanned PDF or textboxes)", file=sys.stderr)
             
        doc.save(output_path)
        print(f"DEBUG: Saved to {output_path}", file=sys.stderr)
        
        return {
            "success": True,
            "output_path": output_path,
            "file_type": "docx",
            "translated_count": translated_count
        }
        
    except Exception as e:
        print(f"DEBUG: Error in translate_docx_file: {str(e)}", file=sys.stderr)
        return {
            "success": False,
            "error": f"DOCX translation failed: {str(e)}"
        }

def create_simple_translated_docx(text: str, output_path: str, source_lang: str, target_lang: str) -> dict:
    """
    Create a professional-looking DOCX with translated text.
    Includes heading detection, bold term preservation, and proper styling.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import re
        
        doc = Document()
        
        # Set up document margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # Clean up page markers like "--- Page 1 ---"
        text = re.sub(r'---\s*Page\s*\d+\s*---\n*', '\n\n[PAGE_BREAK]\n\n', text)
        
        # Replace column separators with section marker
        text = re.sub(r'\n*---\n*', '\n\n[SECTION_BREAK]\n\n', text)
        
        # Split by double newlines to paragraphs
        paragraphs = text.split('\n\n')
        
        # Remove consecutive duplicate paragraphs (common from PDF extraction issues)
        deduped_paragraphs = []
        seen_recent = set()  # Track recently seen paragraph texts within a sliding window
        for para in paragraphs:
            stripped = para.strip()
            if not stripped:
                deduped_paragraphs.append(para)
                continue
            # Skip special markers from dedup check
            if stripped in ('[PAGE_BREAK]', '[SECTION_BREAK]'):
                deduped_paragraphs.append(para)
                seen_recent.clear()  # Reset on page/section breaks
                continue
            # Normalize whitespace for comparison
            normalized = ' '.join(stripped.split())
            if normalized not in seen_recent:
                deduped_paragraphs.append(para)
                seen_recent.add(normalized)
            else:
                print(f"DEBUG: Removing duplicate paragraph: '{stripped[:80]}...'", file=sys.stderr)
        
        paragraphs = deduped_paragraphs
        
        print(f"DEBUG: Creating professional DOCX - {len(paragraphs)} sections (after dedup)", file=sys.stderr)
        
        translated_count = 0
        
        for i, para in enumerate(paragraphs):
            para = para.strip()
            if not para:
                continue
            
            # Handle special markers
            if para == '[PAGE_BREAK]':
                # Add page break
                doc.add_page_break()
                continue
            elif para == '[SECTION_BREAK]':
                # Add a horizontal rule / section separator
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run('─' * 50)
                run.font.color.rgb = RGBColor(150, 150, 150)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                continue
            
            # Translate the paragraph
            result = translate_long_text(para, source_lang, target_lang)
            if result["success"]:
                translated_text = result["translated_text"]
                translated_count += 1
            else:
                print(f"DEBUG: Para {i} translation failed: {result.get('error')}", file=sys.stderr)
                translated_text = para  # Keep original if failed
            
            # Detect if this is a heading (short text, likely a title)
            is_heading = (
                len(translated_text) < 100 and 
                '\n' not in translated_text and
                not translated_text.endswith('.') and
                not translated_text.endswith(',')
            )
            
            # Detect centered/special text (often short lines with dashes or special markers)
            is_centered = (
                translated_text.startswith('-') and translated_text.endswith('-') or
                translated_text.startswith('—') and translated_text.endswith('—') or
                'hereinafter' in translated_text.lower() or
                'referred to as' in translated_text.lower() or
                'nachstehend' in para.lower()  # German "hereinafter"
            )
            
            if is_heading:
                # Add as heading
                heading = doc.add_heading(translated_text, level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif is_centered:
                # Add as centered paragraph
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(translated_text)
                run.italic = True
            else:
                # Regular paragraph - detect and bold quoted terms
                p = doc.add_paragraph()
                
                # Find terms in quotes like "PVA" or "Solar Park" and make them bold
                pattern = r'[""„]([^""„"]+)["""]'
                last_end = 0
                
                for match in re.finditer(pattern, translated_text):
                    # Add text before the match
                    if match.start() > last_end:
                        p.add_run(translated_text[last_end:match.start()])
                    
                    # Add the quoted term with quotes and bold
                    quoted_run = p.add_run(f'"{match.group(1)}"')
                    quoted_run.bold = True
                    
                    last_end = match.end()
                
                # Add remaining text after the last quoted term
                if last_end > 0 and last_end < len(translated_text):
                    p.add_run(translated_text[last_end:])
                elif last_end == 0:
                    # No quoted terms were found - add the full text as a single run
                    p.add_run(translated_text)
            
            # Add some spacing between paragraphs
            if not is_heading:
                p = doc.paragraphs[-1]
                p.paragraph_format.space_after = Pt(8)
        
        doc.save(output_path)
        print(f"DEBUG: Professional DOCX created with {translated_count} translated sections", file=sys.stderr)
        
        return {"success": True, "output_path": output_path, "translated_count": translated_count}
    except Exception as e:
        return {"success": False, "error": f"Fallback creation failed: {str(e)}"}

def convert_pdf_to_docx(pdf_path: str, docx_path: str) -> dict:
    """
    Convert PDF to DOCX using pdf2docx.
    Note: This feature requires pdf2docx which is not available on Render free tier.
    """
    try:
        from pdf2docx import Converter
        
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        
        return {"success": True, "output_path": docx_path}
    except ImportError:
        return {"success": False, "error": "PDF to DOCX conversion not available (pdf2docx not installed)"}
    except Exception as e:
        return {"success": False, "error": f"PDF conversion failed: {str(e)}"}

def translate_document_file(file_path: str, output_path: str, source_lang: str, target_lang: str) -> dict:
    """
    Translate document preserving format (converts PDF to DOCX).
    """
    print(f"DEBUG: translate_document_file called with {file_path}", file=sys.stderr)
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.docx':
        return translate_docx_file(file_path, output_path, source_lang, target_lang)
    
    elif ext == '.pdf':
        # For PDFs: Extract text directly and create translated DOCX
        # This is more reliable than pdf2docx conversion
        print(f"DEBUG: Processing PDF file...", file=sys.stderr)
        
        if not output_path.endswith('.docx'):
            output_path += '.docx'
        
        # Step 1: Extract text from PDF using pdfplumber
        extraction = extract_text_from_pdf(file_path)
        
        if not extraction["success"]:
            print(f"DEBUG: PDF text extraction failed: {extraction.get('error')}", file=sys.stderr)
            return {"success": False, "error": f"Failed to extract text from PDF: {extraction.get('error')}"}
        
        text = extraction.get("text", "")
        page_count = extraction.get("page_count", 0)
        
        print(f"DEBUG: Extracted {len(text)} characters from {page_count} pages", file=sys.stderr)
        
        if not text or not text.strip():
            return {"success": False, "error": "No text found in PDF. The PDF might be scanned/image-based."}
        
        # Step 2: Create translated DOCX
        result = create_simple_translated_docx(text, output_path, source_lang, target_lang)
        
        if result["success"]:
            result["page_count"] = page_count
            result["original_chars"] = len(text)
            print(f"DEBUG: PDF translation complete. Output: {output_path}", file=sys.stderr)
        
        return result
        
    else:
        return {"success": False, "error": "Unsupported file type for format preservation"}

# Quick test when run directly
if __name__ == "__main__":
    print("Document Processor Module")
    print("Supported formats: PDF, DOCX, TXT")
