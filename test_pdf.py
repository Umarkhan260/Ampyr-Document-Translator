"""
Test PDF to DOCX conversion and translation
"""
import sys
sys.path.insert(0, '.')

from document_processor import convert_pdf_to_docx, translate_docx_file, extract_text_from_pdf
from docx import Document
import os

def test_pdf_conversion(pdf_path):
    print(f"\n=== Testing PDF: {pdf_path} ===\n")
    
    # Step 1: Extract text from PDF
    print("1. Extracting text from PDF...")
    extraction = extract_text_from_pdf(pdf_path)
    if extraction["success"]:
        text = extraction["text"]
        print(f"   Extracted {len(text)} characters from {extraction.get('page_count', '?')} pages")
        print(f"   First 500 chars: {text[:500]}...")
    else:
        print(f"   FAILED: {extraction['error']}")
        return
    
    # Step 2: Convert PDF to DOCX
    print("\n2. Converting PDF to DOCX...")
    temp_docx = pdf_path + ".test.docx"
    result = convert_pdf_to_docx(pdf_path, temp_docx)
    if result["success"]:
        print(f"   Converted successfully to: {temp_docx}")
    else:
        print(f"   FAILED: {result['error']}")
        return
    
    # Step 3: Analyze the DOCX structure
    print("\n3. Analyzing DOCX structure...")
    doc = Document(temp_docx)
    print(f"   Paragraphs: {len(doc.paragraphs)}")
    print(f"   Tables: {len(doc.tables)}")
    
    # Check paragraphs with text
    paras_with_text = [p for p in doc.paragraphs if p.text.strip()]
    print(f"   Paragraphs with text: {len(paras_with_text)}")
    
    # Check runs in paragraphs
    total_runs = 0
    runs_with_text = 0
    for para in doc.paragraphs:
        for run in para.runs:
            total_runs += 1
            if run.text and run.text.strip():
                runs_with_text += 1
    
    print(f"   Total runs: {total_runs}")
    print(f"   Runs with text: {runs_with_text}")
    
    # Show first few paragraphs
    print("\n4. First 5 paragraphs content:")
    for i, para in enumerate(doc.paragraphs[:5]):
        print(f"   Para {i}: '{para.text[:100]}...' (runs: {len(para.runs)})")
        for j, run in enumerate(para.runs[:3]):
            print(f"      Run {j}: '{run.text[:50]}...'")
    
    # Cleanup
    if os.path.exists(temp_docx):
        os.remove(temp_docx)
    
    print("\n=== Test Complete ===")

if __name__ == "__main__":
    # Use command line argument or default test file
    import sys
    if len(sys.argv) > 1:
        test_pdf_conversion(sys.argv[1])
    else:
        print("Usage: python test_pdf.py <path_to_pdf>")
        print("No PDF specified, looking for test files...")
        if os.path.exists("test_input.pdf"):
            test_pdf_conversion("test_input.pdf")
