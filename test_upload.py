from docx import Document
import requests
import os

# Create a test DOCX
doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is a test paragraph for translation.')
doc.add_paragraph('Another paragraph with more text.')
doc.save('test_input.docx')

print("Created test_input.docx")

# Upload
url = 'http://127.0.0.1:5000/upload-document'
files = {'file': open('test_input.docx', 'rb')}
data = {'source_lang': 'en', 'target_lang': 'es', 'summarize': 'false'}

print("Uploading...")
try:
    response = requests.post(url, files=files, data=data)
    
    if response.status_code == 200:
        with open('test_output.docx', 'wb') as f:
            f.write(response.content)
        print("Downloaded test_output.docx")
        
        # Verify content
        doc_out = Document('test_output.docx')
        print("Output paragraphs:")
        for p in doc_out.paragraphs:
            print(f"- {p.text}")
            
    else:
        print(f"Error: {response.status_code} - {response.text}")
        
except Exception as e:
    print(f"Request failed: {e}")
